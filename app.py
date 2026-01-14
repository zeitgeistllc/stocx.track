import streamlit as st
from docx import Document
from datetime import date
import io

st.set_page_config(page_title="Zeitgeist LLC – Schwab Docs", layout="centered")

st.title("Zeitgeist LLC – Schwab Transfer Documents")
st.caption("Generate required Word documents for Charles Schwab")

COMPANY_NAME = "Zeitgeist LLC"
OWNER_NAME = "Niv Chen"
CONFIRMATION_NUMBER = "TA04557418"

st.header("Account Information")

delivering_firm = st.text_input("Delivering Firm Name")
delivering_account = st.text_input("Delivering Account Number")
schwab_account = st.text_input("Schwab Account Number (optional)")
execution_date = st.date_input("Execution Date", value=date.today())

def generate_corporate_resolution():
    doc = Document()
    doc.add_heading(f"{COMPANY_NAME}\nCORPORATE RESOLUTION", level=1)
    doc.add_paragraph(
        f"The undersigned, being the sole Member and Manager of {COMPANY_NAME}, "
        "hereby adopts the following resolutions.\n\n"
        f"{OWNER_NAME} is the sole Member and Manager and owns 100% of the company.\n\n"
        f"{OWNER_NAME} is authorized to act on behalf of the company with Charles Schwab.\n\n"
        f"Executed on {execution_date}."
    )
    doc.add_paragraph(f"{OWNER_NAME}\nSole Member and Manager\n{COMPANY_NAME}")
    return doc

def generate_relinquishment_letter():
    doc = Document()
    doc.add_heading("Letter to Relinquish Account Ownership", level=1)
    doc.add_paragraph(
        f"Re: Confirmation {CONFIRMATION_NUMBER}\n\n"
        f"I, {OWNER_NAME}, relinquish ownership of account {delivering_account} "
        f"at {delivering_firm} for transfer to Schwab."
    )
    doc.add_paragraph(f"{OWNER_NAME}\n{COMPANY_NAME}")
    return doc

if st.button("Generate Documents"):
    if delivering_firm and delivering_account:
        res = generate_corporate_resolution()
        rel = generate_relinquishment_letter()

        res_buf, rel_buf = io.BytesIO(), io.BytesIO()
        res.save(res_buf)
        rel.save(rel_buf)

        st.download_button("Download Corporate Resolution", res_buf.getvalue(),
            "Zeitgeist_LLC_Corporate_Resolution.docx")
        st.download_button("Download Relinquishment Letter", rel_buf.getvalue(),
            "Zeitgeist_LLC_Letter_to_Relinquish_Account_Ownership.docx")
    else:
        st.error("Delivering firm and account number required")
